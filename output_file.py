from docx.enum.text import WD_COLOR_INDEX
import re
import os
import  moce_dir
###that part is for bold words in txt
##convert txt to docx and bold the word
##there is few problem with that code.


curr_directory=os.getcwd()
curr_directory_from = os.getcwd() + "\\takefrom"
curr_directory_to = os.getcwd() + "\\taketo"


def output_File(filepath,tokens):
    if tokens is None:
        print("tokens = None")
    print('tokens: ',tokens)

    # print('here: ',filepath)
    path = filepath.split(os.sep)
    path = path[1].split(".")
    new_path=path[0]
    print('new_path: ',new_path)

    path=curr_directory+"\\"+filepath

    document1 = Document()
    # document1.add_heading(new_path, 0)
    myfile = open(path).read()
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c', ' ', myfile)  # remove all non-XML-compatible characters
    document1.add_paragraph(myfile)
    document1.save(curr_directory_to + "\\" + new_path + '.docx')

    doc = Document(curr_directory_to+"\\"+ new_path + '.docx')
    new_document = Document()
    new_document.add_heading(new_path,0)
    p2 = new_document.add_paragraph()

    for t in tokens:
        for paragraph in doc.paragraphs:
                if t in paragraph.text:
                    substrings = paragraph.text.split(t)
                    for substring in substrings[:-1]:
                        p2.add_run(substring)
                        font = p2.add_run(t).font
                        font.highlight_color = WD_COLOR_INDEX.YELLOW
                    p2.add_run(substrings[-1])
                else:
                    p2.add_run(paragraph.text)


    new_document.save(curr_directory_from + "\\" + new_path + '.docx')
    os.startfile(curr_directory_from+"\\"+new_path +'.docx')
    moce_dir.delete_taketo()


if __name__ == "__main__":
    tokens2='sleep','poppy'

    output_File("storage\\101.txt", tokens2)
    # direct=os.listdir(curr_directory_from)
    # for i in direct:
    #     document = Document()
    #     document.add_heading(i, 0)
    #     myfile = open(curr_directory_from+"\\"+ i).read()
    #     myfile = re.sub(r'[^\x00-\x7F]+|\x0c', ' ', myfile)  # remove all non-XML-compatible characters
    #     p = document.add_paragraph(myfile)
    #     document.save(curr_directory_to +"\\"+ i + '.docx')
    #
    #
    # doc = Document(curr_directory_to+"\\"+"100.txt.docx")
    # new_doc = Document()
    # #print('word:',word)
    #
    # p2 = new_doc.add_paragraph()
    # for paragraph in doc.paragraphs:
    #     if 'sleep' in paragraph.text:
    #         substrings = paragraph.text.split('sleep')
    #         for substring in substrings[:-1]:
    #             p2.add_run(substring)
    #             font = p2.add_run('sleep').font
    #             font.highlight_color = WD_COLOR_INDEX.YELLOW
    #         p2.add_run(substrings[-1])
    #     else:
    #         p2.add_run(paragraph.text)
    #
    # new_doc.save('t1.docx')